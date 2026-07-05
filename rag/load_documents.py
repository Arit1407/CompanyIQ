"""
Load Annual Reports (DOCX + PDF)

CompanyIQ RAG Pipeline
"""

from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader

from config.config import REPORTS_PATH


# ===================================================
# Load DOCX
# ===================================================

def load_docx(file_path: Path, company_name: str):

    doc = DocxDocument(file_path)

    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    merged_text = "\n".join(text)

    return [
        Document(
            page_content=merged_text,
            metadata={
                "company": company_name,
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx"
            }
        )
    ]


# ===================================================
# Load PDF (Merge all pages into ONE document)
# ===================================================

def load_pdf(file_path: Path, company_name: str):

    loader = PyPDFLoader(str(file_path))
    pages = loader.load()

    full_text = []

    total_pages = len(pages)

    for page in pages:

        if page.page_content.strip():
            full_text.append(page.page_content)

    merged_text = "\n\n".join(full_text)

    return [
        Document(
            page_content=merged_text,
            metadata={
                "company": company_name,
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "pdf",
                "total_pages": total_pages
            }
        )
    ]


# ===================================================
# Load All Reports
# ===================================================

def load_documents():

    all_documents = []

    report_root = Path(REPORTS_PATH)

    if not report_root.exists():
        raise FileNotFoundError(
            f"Reports folder not found: {REPORTS_PATH}"
        )

    for company_folder in sorted(report_root.iterdir()):

        if not company_folder.is_dir():
            continue

        company_name = company_folder.name

        print(f"Loading {company_name}")

        for file in sorted(company_folder.iterdir()):

            try:

                if file.suffix.lower() == ".docx":

                    all_documents.extend(
                        load_docx(file, company_name)
                    )

                elif file.suffix.lower() == ".pdf":

                    all_documents.extend(
                        load_pdf(file, company_name)
                    )

            except Exception as e:

                print(f"Failed to load {file.name}")
                print(e)

    print("\n===================================")
    print("Documents Loaded")
    print("===================================")
    print(f"Total Documents : {len(all_documents)}")

    return all_documents


# ===================================================
# Test
# ===================================================

if __name__ == "__main__":

    documents = load_documents()

    print("\n===================================")
    print("Sample Metadata")
    print("===================================")

    print(documents[0].metadata)

    print("\n===================================")
    print("First 500 Characters")
    print("===================================\n")

    print(documents[0].page_content[:500])